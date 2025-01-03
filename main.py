from fastapi import FastAPI
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from conexion import database as connection
from modelos import *
from limitaciones import *
from fastapi import HTTPException
from typing import List
from fastapi import Query, UploadFile, File
import pandas as pd 
import numpy as np
import pytesseract as tess
from PIL import Image
import cv2
import re
import json
from io import BytesIO 
import os
import io
from docx import Document
from fastapi.responses import StreamingResponse
os.environ['TESSDATA_PREFIX'] = r'C:\Program Files\Tesseract-OCR\tessdata'


app = FastAPI(title='api ccc    ',
              description='api para maneacion de archivos',#aqui se nombra y crea la app que vas a subir en el servidor 
              version='1.1')#tambien se ocupa el nombre de este archivo 

'''origins = [#si se quiere dar acceso desde un servidor exacto
    "http://localhost",
    "http://localhost:80",
    "http://127.0.0.1",
    "http://127.0.0.1/front/index.html:80"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,  # Lista de orígenes permitidos
    allow_credentials=True,
    allow_methods=["*"],  # Permite todos los métodos
    allow_headers=["*"],  # Permite todos los encabezados
)
'''

app.add_middleware( #si que quiere permisitr acceso desde cualquier front
    CORSMiddleware,
    allow_origins=["*"],  # Permite todas las solicitudes de cualquier origen
    allow_credentials=True,
    allow_methods=["*"],  # Permite todos los métodos HTTP
    allow_headers=["*"],  # Permite todos los encabezados
)

tess.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

@app.on_event('startup')#aqui los procesos que se general al arranque 
async def startup_event():
    if connection.is_closed():
        connection.connect()#con esto se crean tablas 
        connection.create_tables([user,auto,cliente,facturas])

@app.on_event('shutdown')
def shutdown_event():
    if not connection.is_closed():
        connection.close()
@app.get('/')
async def index():
    return 'hola dfrr'

def Auto():
    @app.post('/auto')#agreagar auto 
    async def create_auto(auto_request: AutoRequestModel):
        new_auto =auto.create(
                auto=auto_request.auto,
                modelo=auto_request.modelo,
                version=auto_request.version,
                color=auto_request.color,
                precio=auto_request.precio,
                transmision=auto_request.transmision,
                motor=auto_request.motor,
                n_motor=auto_request.n_motor,
                n_puertas=auto_request.n_puertas,
                tipo=auto_request.tipo,
                chasis=auto_request.chasis,
                fecha_aprovacion=auto_request.fecha_aprovacion,
                certificado=auto_request.certificado
        )
        return auto_request
    
    
    @app.get('/auto/{Auto}')#para un auto en especifico 
    async def get_auto_data(Auto, page: int = 0, limit: int = 0):
        get_auto1 = auto.select().where(auto.auto == Auto).first()
        if get_auto1:
            return AutoResponseModel(
                referencia=get_auto1.referencia, 
                auto=get_auto1.auto,
                modelo=get_auto1.modelo,
                version=get_auto1.version,
                color=get_auto1.color,
                precio=get_auto1.precio,
                transmision=get_auto1.transmision,
                motor=get_auto1.motor,
                n_motor=get_auto1.n_motor,
                n_puertas=get_auto1.n_puertas,
                tipo=get_auto1.tipo,
                chasis=get_auto1.chasis,
                fecha_aprovacion=get_auto1.fecha_aprovacion,
                certificado=get_auto1.certificado
            )
        else:
            raise HTTPException(404, 'no encontrado')
    
    @app.get('/auto1/', response_model=List[AutoResponseModel])#para todos los autos en general
    async def get_all_autos(page: int = 0, limit: int = 0):
        autos2 = auto.select()

        if autos2:
            resultados = []
            for auto_item in autos2:
                auto_dict = {
                    'referencia': auto_item.referencia,
                    'auto': auto_item.auto,
                    'modelo': auto_item.modelo,
                    'version': auto_item.version,
                    'color': auto_item.color,
                    'precio': auto_item.precio,
                    'transmision': auto_item.transmision,
                    'motor': auto_item.motor,
                    'n_motor': auto_item.n_motor,
                    'n_puertas': auto_item.n_puertas,
                    'tipo': auto_item.tipo,
                    'chasis': auto_item.chasis,
                    'fecha_aprovacion': auto_item.fecha_aprovacion,
                    'certificado': auto_item.certificado
                }
                resultados.append(auto_dict)
            return resultados
        else:
            raise HTTPException(status_code=404, detail='No se encontraron autos')

    @app.get('/auto2/', response_model=List[AutoResponseModel])#para todos los autos con un nombre
    async def get_all_autos(auto_name: str = Query(..., description="Nombre del coche a buscar")):
        autos_filtered = auto.select().where(auto.auto == auto_name)

        if autos_filtered:
            resultados = []
            for auto_item in autos_filtered:
                auto_dict = {
                    'referencia': auto_item.referencia,
                    'auto': auto_item.auto,
                    'modelo': auto_item.modelo,
                    'version': auto_item.version,
                    'color': auto_item.color,
                    'precio': auto_item.precio,
                    'transmision': auto_item.transmision,
                    'motor': auto_item.motor,
                    'n_motor': auto_item.n_motor,
                    'n_puertas': auto_item.n_puertas,
                    'tipo': auto_item.tipo,
                    'chasis': auto_item.chasis,
                    'fecha_aprovacion': auto_item.fecha_aprovacion,
                    'certificado': auto_item.certificado
                }
                resultados.append(auto_dict)
            return resultados
        else:
            raise HTTPException(status_code=404, detail=f'No se encontraron autos con el nombre "{auto_name}"')
     
     
    @app.put("/auto/{referencia}", response_model=AutoResponseModel)
    async def update_auto(referencia: int, auto_request: AutoRequestModel):
        # Intentar obtener el auto existente por la referencia
        auto_obj = auto.get_or_none(auto.referencia == referencia)
        if not auto_obj:
            raise HTTPException(status_code=404, detail="Auto no encontrado")

        # Actualizar los campos del auto con los datos proporcionados en auto_request
        auto_obj.auto = auto_request.auto
        auto_obj.modelo = auto_request.modelo
        auto_obj.version = auto_request.version
        auto_obj.color = auto_request.color
        auto_obj.precio = auto_request.precio
        auto_obj.transmision = auto_request.transmision
        auto_obj.motor = auto_request.motor
        auto_obj.n_motor = auto_request.n_motor
        auto_obj.n_puertas = auto_request.n_puertas
        auto_obj.tipo = auto_request.tipo
        auto_obj.chasis = auto_request.chasis
        auto_obj.fecha_aprovacion = auto_request.fecha_aprovacion
        auto_obj.certificado = auto_request.certificado

        # Guardar los cambios en la base de datos
        auto_obj.save()

        # Devolver el auto actualizado
        return AutoResponseModel(
            referencia=auto_obj.referencia, 
            auto=auto_obj.auto,
            modelo=auto_obj.modelo,
            version=auto_obj.version,
            color=auto_obj.color,
            precio=auto_obj.precio,
            transmision=auto_obj.transmision,
            motor=auto_obj.motor,
            n_motor=auto_obj.n_motor,
            n_puertas=auto_obj.n_puertas,
            tipo=auto_obj.tipo,
            chasis=auto_obj.chasis,
            fecha_aprovacion=auto_obj.fecha_aprovacion,
            certificado=auto_obj.certificado
        )
 
            
    @app.delete('/auto/{referencia}')
    async def delete_user(referencia):
        delete_auto=auto.select().where(auto.referencia == referencia).first()
        if delete_auto:
            delete_auto.delete_instance()
            return True
        else:
            return HTTPException(404, 'no encontrado')    
            
def clientes():
   @app.post("/importar_clientes/", response_model=List[ClienteResponseModel])
   async def importar_clientes(archivo_excel: UploadFile = File(...)):
        if not archivo_excel.filename.endswith('.xlsx'):
            raise HTTPException(status_code=400, detail="El archivo debe ser un archivo .xlsx")

        try:
            df = pd.read_excel(archivo_excel.file, engine='openpyxl')
            columnas_esperadas = ["nombre", "numero", "direccion", "colonia", "estado", "cp", "pais", "numero_pagos", "orden_envio"]
            if not all(columna in df.columns for columna in columnas_esperadas):
                raise HTTPException(status_code=400, detail="El archivo Excel no contiene todas las columnas necesarias.")
            df['numero'] = df['numero'].astype(str)
            clientes_creados = []
            for _, row in df.iterrows():
                cliente_data = ClienteRequestModel(
                    nombre=row['nombre'],
                    numero=row['numero'],
                    direccion=row['direccion'],
                    colonia=row['colonia'],
                    estado=row['estado'],
                    cp=row['cp'],
                    pais=row['pais'],
                    numero_pagos=row['numero_pagos'],
                    orden_envio=row['orden_envio']
                )
                cliente_nuevo = cliente.create(**cliente_data.dict())
                clientes_creados.append(cliente_nuevo)

            # Usar model_validate en lugar de from_orm
            response_models = [ClienteResponseModel.model_validate(c) for c in clientes_creados]
            return response_models

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error al procesar el archivo: {str(e)}")
        
   @app.post("/enviardatos_cliente/")
   async def enviar_datos_cliente(archivo: UploadFile = File(...)):
        if not archivo.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            raise HTTPException(status_code=400, detail="El archivo debe ser una imagen PNG, JPG o JPEG")

        contenido_imagen = await archivo.read()
        imagen = Image.open(BytesIO(contenido_imagen))
        imagen_np = np.array(imagen)
        imagen_cv2 = cv2.cvtColor(imagen_np, cv2.COLOR_RGB2BGR)

        texto = tess.image_to_string(imagen_cv2)

        datos_cliente = extract_data(texto)
        return JSONResponse(content=datos_cliente)

   def extract_data(text):
            data = {}
            nombre_match = re.search(r'(?<=NOMBRE\s)[\s\S]*?(?=DOMICILIO)', text)
            if nombre_match:
                data['nombre'] = re.sub(r'\s+', ' ', nombre_match.group().strip())

            domicilio_pattern = re.compile(r'DOMICILIO\.\s*([\s\S]*?)COL\.\s*([\s\S]*?)C\.P\.\s*(\d+)')
            domicilio_match = domicilio_pattern.search(text)
            if domicilio_match:
                data['direccion'] = domicilio_match.group(1).strip()
                data['colonia'] = re.sub(r'\s+', ' ', domicilio_match.group(2).strip())
                data['cp'] = int(domicilio_match.group(3).strip())

            estado_match = re.search(r'C\.P\.\s+\d+,\s+([\s\S]*?)(?=,|\n)', text)
            if estado_match:
                data['estado'] = re.sub(r'\s+', ' ', estado_match.group(1).strip())

            data['pais'] = 'México'

            return data
    
    
   @app.post("/agregar_cliente")
   async def agregar_cliente(cliente_data: ClienteRequestModel):
        # Verificar si el cliente con el mismo ID o número ya existe
        cliente_existente = cliente.get_or_none(cliente.id_cliente == cliente_data.id_cliente)
        if cliente_existente:
            raise HTTPException(status_code=400, detail="Cliente con este ID ya existe")

        cliente_existente_numero = cliente.get_or_none(cliente.numero == cliente_data.numero)
        if cliente_existente_numero:
            raise HTTPException(status_code=400, detail="Cliente con este número ya existe")

        # Crear un nuevo cliente en la base de datos
        nuevo_cliente = cliente.create(
            nombre=cliente_data.nombre,
            numero=cliente_data.numero,
            direccion=cliente_data.direccion,
            colonia=cliente_data.colonia,
            estado=cliente_data.estado,
            cp=cliente_data.cp,
            pais=cliente_data.pais,
            numero_pagos=cliente_data.numero_pagos,
            orden_envio=cliente_data.orden_envio
        )

        # Guardar el cliente en la base de datos
        nuevo_cliente.save()

        # Devolver el cliente agregado, utilizando el modelo de respuesta
        return ClienteResponseModel.from_orm(nuevo_cliente)
    
   
            
            
        
   @app.get("/clientes/", response_model=List[ClienteResponseModel])
   async def obtener_todos_los_clientes():
        clientes = cliente.select()
        return [ClienteResponseModel.model_validate(c) for c in clientes]
    
   @app.put("/clientes/{nombre}", response_model=ClienteResponseModel)
   async def actualizar_cliente(nombre: str, cliente_data: ClienteRequestModel):
        cliente_obj = cliente.get_or_none(cliente.nombre == nombre)
        if not cliente_obj:
            raise HTTPException(status_code=404, detail="Cliente no encontrado")
        for key, value in cliente_data.dict().items():
            setattr(cliente_obj, key, value)
        cliente_obj.save()
        return ClienteResponseModel.model_validate(cliente_obj)

   @app.delete("/clientes/{numero}")
   async def eliminar_cliente(numero: int):
        cliente_obj = cliente.get_or_none(cliente.numero == numero)
        if not cliente_obj:
            raise HTTPException(status_code=404, detail="Cliente no encontrado")
        cliente_obj.delete_instance()
        return {"message": "Cliente eliminado correctamente"}



def facturacion():
    
    @app.post('/generar_factura')
    async def create_factura(factura_request: FacturaRequestModel):
        data = factura_request.dict()

        # Cargar el documento de plantilla
        template_path = "venv/plantilla_factura.docx"
        doc = Document(template_path)
        
        def replace_text(doc, placeholder, replacement):
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(replacement))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(replacement))

        # Reemplazar los marcadores de posición con los valores del JSON
        replace_text(doc, "{{id_factura}}", data["id_factura"])
        replace_text(doc, "{{referencia_id}}", data["referencia_id"])
        replace_text(doc, "{{auto}}", data["auto"])
        replace_text(doc, "{{modelo}}", data["modelo"])
        replace_text(doc, "{{version}}", data["version"])
        replace_text(doc, "{{color}}", data["color"])
        replace_text(doc, "{{precio}}", data["precio"])
        replace_text(doc, "{{transmision}}", data["transmision"])
        replace_text(doc, "{{motor}}", data["motor"])
        replace_text(doc, "{{n_motor}}", data["n_motor"])
        replace_text(doc, "{{n_puertas}}", data["n_puertas"])
        replace_text(doc, "{{tipo}}", data["tipo"])
        replace_text(doc, "{{chasis}}", data["chasis"])
        replace_text(doc, "{{fecha_aprovacion}}", data["fecha_aprovacion"])
        replace_text(doc, "{{certificado}}", data["certificado"])
        replace_text(doc, "{{id_cliente_id}}", data["id_cliente_id"])
        replace_text(doc, "{{nombre}}", data["nombre"])
        replace_text(doc, "{{numero}}", data["numero"])
        replace_text(doc, "{{Direccion}}", data["Direccion"])
        replace_text(doc, "{{Colonia}}", data["Colonia"])
        replace_text(doc, "{{Estado}}", data["Estado"])
        replace_text(doc, "{{CP}}", data["CP"])
        replace_text(doc, "{{Pais}}", data["Pais"])
        replace_text(doc, "{{numero_pagos}}", data["numero_pagos"])
        replace_text(doc, "{{orden_envio}}", data["orden_envio"])
        replace_text(doc, "{{fecha}}", data["fecha"])
        replace_text(doc, "{{hora}}", data["hora"])

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return StreamingResponse(file_stream, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={"Content-Disposition": "attachment; filename=factura_generada.docx"})
    @app.post('/factura')
    async def create_factura(factura_request: FacturaRequestModel):
        new_factura = facturas.create(**factura_request.dict())
        return FacturaResponseModel.from_orm(new_factura)
    
    @app.get("/factura/", response_model=List[FacturaResponseModel])
    async def get_all_facturas():
        all_facturas = facturas.select()
        return [FacturaResponseModel.from_orm(factura) for factura in all_facturas]


    @app.put('/factura/{factura_id}', response_model=FacturaResponseModel)
    async def update_factura(factura_id: int, factura_request: FacturaRequestModel):
        # Buscar la factura existente usando el ID
        factura_obj = facturas.get_or_none(facturas.id_factura == factura_id)
        if not factura_obj:
            raise HTTPException(status_code=404, detail="Factura no encontrada")

        # Actualizar los campos de la factura con los datos proporcionados en factura_request
        for field, value in factura_request.dict(exclude_unset=True).items():
            setattr(factura_obj, field, value)

        # Guardar los cambios en la base de datos
        factura_obj.save()

        # Devolver la factura actualizada como respuesta
        return FacturaResponseModel.from_orm(factura_obj)


    @app.delete('/factura/{factura_id}')
    async def delete_factura(factura_id: int):
        factura_obj = facturas.get_or_none(facturas.id_factura == factura_id)
        if not factura_obj:
            raise HTTPException(status_code=404, detail="Factura no encontrada")
        factura_obj.delete_instance()
        return {"message": "Factura eliminada correctamente"}




                

def usuario():
    @app.post('/user')
    async def create_user(user_request: UserRequestModel):
        new_user =user.create(
            nombre=user_request.nombre,
            password=user_request.password
        )
        return user_request


    @app.get('/user/{nombre}/{password}')
    async def get_user(nombre: str, password: str):   
        get_nombre = user.select().where((user.nombre == nombre) & (user.password == password)).first()
        if get_nombre:
            return {"status": "success", "message": "Login exitoso"}  # Devolver una respuesta JSON
        else:
            raise HTTPException(status_code=404, detail='No encontrado')



    @app.get('/user/{user_id}')
    async def get_user(user_id, page: int= 0, limit:int=0):
        get_nombre=user.select().where(user.id == user_id).first()
        if get_nombre:
            return UserResponseModel(id=get_nombre.id, nombre=get_nombre.nombre, password=get_nombre.password)
        else:
            return HTTPException(404, 'no encontrado')
        
    @app.put('/user/{user_id}')
    async def update_user(user_id: int, user_request: UserRequestModel):
        updated_user = user.select().where(user.id == user_id).first()
        if updated_user:
            updated_user.nombre = user_request.nombre
            updated_user.password = user_request.password
            updated_user.save()  # Guardar los cambios en la base de datos
            return {"message": "Usuario actualizado correctamente"}
        else:
            raise HTTPException(status_code=404, detail="Usuario no encontrado")

    @app.delete('/user/{user_id}')
    async def delete_user(user_id):
        delete_nombre=user.select().where(user.id == user_id).first()
        if delete_nombre:
            delete_nombre.delete_instance()
            return True
        else:
            return HTTPException(404, 'no encontrado')    
        
usuario()
Auto()
clientes()
facturacion()